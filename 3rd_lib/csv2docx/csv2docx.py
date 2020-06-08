import argparse
import csv
import sys

from docx import *
from docx import shared
from docx.oxml.shared import OxmlElement, qn


class DataCSV:
    def __init__(self, file):
        if file.split('.')[-1] == "csv":
            self.file = file
        else:
            self.file = file + ".csv"

    def operate_file(self):
        steps_result_list = []
        with open(self.file, encoding='utf-8') as f:
            file_csv = csv.DictReader(f)

            for row in file_csv:
                steps_result_list.append(
                    {'title': row["用例标题"], 'steps': row['步骤'],
                     'expect': row['预期']})

        return steps_result_list


class MyDoc:
    def __init__(self):
        self.document = Document()  # use the Document class from docx
        self.set_doc_style()

    def set_doc_style(self):
        style = self.document.styles['Normal']
        font = style.font
        font.name = "等线"
        font.size = shared.Pt(10)

    def add_local_heading(self, content, level=3):
        p = self.document.add_heading(content, level)
        p.add_run().bold = True

    def create_table(self, header, records):
        def shade_cells(cells, shade):
            for cell in cells:
                tcPr = cell._tc.get_or_add_tcPr()
                tcVAlign = OxmlElement("w:shd")
                tcVAlign.set(qn("w:fill"), shade)
                tcPr.append(tcVAlign)

        table = self.document.add_table(rows=2, cols=5)
        table.style = 'TableGrid'

        # merge the first row
        first_row = table.rows[0]
        first_row.cells[0].merge(first_row.cells[-1])
        first_row.cells[0].paragraphs[0].add_run(
            'Objective1: '+ records[-1]).bold = True
        shade_cells([first_row.cells[0]], '#A4D39F')

        # create the header
        hdr_cells = table.rows[1].cells
        for i in range(len(hdr_cells)):
            hdr_cells[i].paragraphs[0].add_run(header[i]).bold = True
            shade_cells([hdr_cells[i]], '#E0E0E0')

        # fill the data to table
        row_cells = table.add_row().cells
        for index in range(len(row_cells)):
            row_cells[index].text = records[index]

    def save_doc(self, doc_name):
        if doc_name.split('.')[-1] == 'docx':
            self.document.save(doc_name)
        else:
            self.document.save(doc_name + '.docx')


def parse_argument(argv):
    parser = argparse.ArgumentParser()
    parser.add_argument('--csv file', type=str,
                        help='please specify file exported from Zentao')
    parser.add_argument('--docx file', type=str,
                        help=' please specify the docx name')
    return parser.parse_args(argv)


if __name__ == "__main__":
    csv_name = sys.argv[1]
    doc_name = sys.argv[2]

    # deal with the csv file
    my_csv = DataCSV(csv_name)
    steps_result = my_csv.operate_file()
    # print(steps_result)

    one_record = []
    for item in steps_result:
        one_record.append(
            [item['steps'], ' ', item['expect'], 'Pass', ' ', item['title']])

    # print(one_record)
    # export data to my doc
    my_doc = MyDoc()
    header = (
        'Description/Action', 'Specific Data/ Condition(s)',
        'Expected Result',
        "Pass/Fail", 'Comments')

    for index, item in enumerate(one_record):
        my_doc.add_local_heading(
            "Test Case {}:{}".format(str(index + 1), item[-1]))
        my_doc.create_table(header, item)

    my_doc.save_doc(doc_name)
